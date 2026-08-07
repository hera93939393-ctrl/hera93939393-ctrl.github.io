import json
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[1]
ROOT = PROJECT_ROOT / "generated" / "KTW VOCA"
SOURCE = PROJECT_ROOT / "data"
CODE_BY_LEVEL = {
    "BASIC": "BA", "PRE-INTERMEDIATE": "PA", "INTERMEDIATE": "IA", "HIGH": "HA",
    "JUPITER": "JUPITER", "SATURN": "SATURN", "URANUS": "URANUS", "NEPTUNE": "NEPTUNE",
}
datasets = {
    "A": json.loads((SOURCE / "selected_A.json").read_text(encoding="utf-8")),
    "B": json.loads((SOURCE / "selected_B.json").read_text(encoding="utf-8")),
}
level_filter = sys.argv[1] if len(sys.argv) > 1 else None
scan_root = ROOT / level_filter if level_filter else ROOT
files = sorted(scan_root.rglob("*.docx"))
expected_file_count = 4 if level_filter else 32
errors = []
for path in files:
    try:
        with ZipFile(path) as zf:
            bad = zf.testzip()
            if bad:
                errors.append(f"{path}: corrupt member {bad}")
        doc = Document(path)
        for table_index, table in enumerate(doc.tables):
            tbl_pr = table._tbl.tblPr
            jc = tbl_pr.find(qn("w:jc"))
            layout = tbl_pr.find(qn("w:tblLayout"))
            grid_total = sum(int(col.get(qn("w:w"), "0")) for col in table._tbl.tblGrid)
            if jc is None or jc.get(qn("w:val")) != "center":
                errors.append(f"{path}: table {table_index} is not centered")
            if layout is None or layout.get(qn("w:type")) != "fixed":
                errors.append(f"{path}: table {table_index} is not fixed width")
            if grid_total > 10450:
                errors.append(f"{path}: table {table_index} grid too wide ({grid_total})")
            for row in table.rows:
                for cell in row.cells:
                    if cell.vertical_alignment != WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                        errors.append(f"{path}: table {table_index} cell not vertically centered")
                        break
                    if any(p.alignment != WD_ALIGN_PARAGRAPH.CENTER for p in cell.paragraphs):
                        errors.append(f"{path}: table {table_index} cell text not centered")
                        break
        if len(doc.tables) != 5:
            errors.append(f"{path}: expected 5 tables, got {len(doc.tables)}")
        exam_tables = [t for t in doc.tables if len(t.rows) == 26 and len(t.columns) == 8]
        if len(exam_tables) != 2:
            errors.append(f"{path}: expected 2 exam tables, got {len(exam_tables)}")
            continue
        for idx, table in enumerate(exam_tables):
            if len(table.rows) != 26 or len(table.columns) != 8:
                errors.append(f"{path}: table {idx} is {len(table.rows)}x{len(table.columns)}")
        all_table_text = " ".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
        if "★ Basic / ★★ Challenge / ★★★ Master" not in all_table_text:
            errors.append(f"{path}: guide box missing")
        if not doc.sections[0].footer.paragraphs[0]._p.xpath('.//w:drawing'):
            errors.append(f"{path}: footer logo missing")

        level_name = path.parents[1].name
        ab = "A" if path.parent.name == "A형" else "B"
        code = CODE_BY_LEVEL[level_name]
        raw = datasets[ab][code]
        expected = [item for _, item in sorted(enumerate(raw), key=lambda x: (int(x[1].get("tier", 1)), x[1]["day"], x[0]))]
        is_answer = "정답지" in path.stem
        records = {}
        for table in exam_tables:
            for row in table.rows[1:]:
                for no_col, star_col, word_col, meaning_col in ((0, 1, 2, 3), (4, 5, 6, 7)):
                    no_text = row.cells[no_col].text.strip()
                    if no_text:
                        records[int(no_text)] = {
                            "stars": row.cells[star_col].text.count("★"),
                            "word": row.cells[word_col].text.strip(),
                            "meaning": row.cells[meaning_col].text.strip(),
                        }
        if len(records) != len(expected):
            errors.append(f"{path}: expected {len(expected)} items, got {len(records)}")
        last_tier = 0
        for no, item in enumerate(expected, 1):
            actual = records.get(no)
            if not actual:
                continue
            tier = int(item.get("tier", 1))
            if tier < last_tier:
                errors.append(f"{path}: tier order drops at {no}")
            last_tier = tier
            if actual["stars"] != tier:
                errors.append(f"{path}: star/source mismatch at {no}")
            if is_answer:
                if actual["word"] != item["word"] or actual["meaning"] != item["meaning"]:
                    errors.append(f"{path}: answer mismatch at {no}")
            elif no % 2:
                if actual["word"] != item["word"] or actual["meaning"]:
                    errors.append(f"{path}: English prompt/blank mismatch at {no}")
            else:
                if actual["word"] != item["meaning"] or actual["meaning"]:
                    errors.append(f"{path}: Korean prompt/blank mismatch at {no}")
    except Exception as exc:
        errors.append(f"{path}: {exc}")

print(f"FILES={len(files)}")
print(f"ERRORS={len(errors)}")
for error in errors:
    print(error)
raise SystemExit(1 if errors or len(files) != expected_file_count else 0)
