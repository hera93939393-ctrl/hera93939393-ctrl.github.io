import json
import math
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DATA = PROJECT_ROOT / "data" / "selected_A.json"
LOGO = PROJECT_ROOT / "assets" / "학원로고.jpg"
OUT = PROJECT_ROOT / "generated" / "BASIC_A_시험지_샘플.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=20, start=30, bottom=20, end=30):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size, bold=False, color="000000", name="함초롬바탕"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def fitted_font_size(text, width_cm, height_cm, max_size=10.5):
    """Choose a conservative size so wrapped text fits the fixed-height cell."""
    if not text:
        return max_size
    visual_units = sum(1.75 if unicodedata.east_asian_width(ch) in ("W", "F", "A") else 1.0 for ch in text)
    usable_width_pt = max(20.0, width_cm * 28.346 - 5.0)
    usable_height_pt = max(10.0, height_cm * 28.346 - 2.0)
    for size in (max_size, 9.5, 8.5, 7.5, 6.5, 6.0):
        units_per_line = max(1.0, usable_width_pt / (size * 0.58))
        lines = max(1, math.ceil(visual_units / units_per_line))
        if lines * size * 0.80 <= usable_height_pt:
            return size
    return 6.0


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def set_table_borders(table, color="000000", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_cm):
    """Lock Word table grid/total width so content cannot expand it past the page."""
    widths_twips = [round(width * 567) for width in widths_cm]
    total_twips = sum(widths_twips)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_twips[min(idx, len(widths_twips) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Cm(widths_cm[min(idx, len(widths_cm) - 1)])


def add_header(doc, level, ab):
    title_table = doc.add_table(rows=1, cols=1)
    title_table.autofit = False
    title_table.alignment = 1
    title_cell = title_table.cell(0, 0)
    title_cell.width = Cm(13.8)
    set_cell_margins(title_cell, 0, 30, 0, 30)
    shade(title_cell, "1F3864")
    title_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(0)
    set_font(title_p.add_run("KTW VOCA MASTER CONTEST"), 18, True, "FFFFFF", "Times New Roman")
    set_table_geometry(title_table, [13.8])

    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(0)
    gap.paragraph_format.line_spacing = Pt(1)

    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    widths_cm = [4.1, 3.1, 10.0, 1.2]
    for cell, width in zip(table.rows[0].cells, widths_cm):
        cell.width = Cm(width)
        set_cell_margins(cell, 0, 25, 0, 25)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    c0, c1, c2, c3 = table.rows[0].cells
    shade(c0, "DDEBF7")
    shade(c1, "DDEBF7")
    shade(c3, "ED7D31")
    set_font(c0.paragraphs[0].add_run("Class :"), 11, False, "000000", "Times New Roman")
    set_font(c1.paragraphs[0].add_run(level), 11, False, "000000", "Times New Roman")
    set_font(c2.paragraphs[0].add_run("Name : __________________________"), 11, False, "000000", "Times New Roman")
    set_font(c3.paragraphs[0].add_run(f"({ab}형)"), 10.5, True, "FFFFFF", "맑은 고딕")
    set_table_geometry(table, widths_cm)
    for source_table in (title_table, table):
     for row in source_table.rows:
      for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        tc_pr.append(borders)
    guide_table = doc.add_table(rows=1, cols=1)
    guide_table.autofit = False
    guide_table.alignment = 1
    guide_cell = guide_table.cell(0, 0)
    guide_cell.width = Cm(18.4)
    guide_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(guide_cell, 20, 30, 20, 30)
    shade(guide_cell, "EAF2F8")
    guide = guide_cell.paragraphs[0]
    guide.alignment = WD_ALIGN_PARAGRAPH.CENTER
    guide.paragraph_format.space_before = Pt(0)
    guide.paragraph_format.space_after = Pt(0)
    guide.paragraph_format.line_spacing = 0.9
    guide_text = "★ Basic / ★★ Challenge / ★★★ Master (쉬운문제부터 차근히 풀어보세요! 끝까지 포기하지 않는것이 가장 중요합니다:))"
    set_font(guide.add_run(guide_text), 9.0, True, "1F3864", "맑은 고딕")
    tc_pr = guide_cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tc_pr.append(borders)
    set_table_geometry(guide_table, [18.4])


def add_exam_table(doc, page_items, start_no, row_height_cm, is_answer=False):
    table = doc.add_table(rows=26, cols=8)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    widths_cm = [0.85, 0.60, 3.45, 4.30] * 2
    widths = [Cm(value) for value in widths_cm]
    headers = ["No", "", "Word", "Meaning", "No", "", "Word", "Meaning"]
    for col, (width, label) in enumerate(zip(widths, headers)):
        cell = table.rows[0].cells[col]
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, "D9E2F3")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), 9.5 if label == "No" else 10.5, True, name="Times New Roman")
    table.rows[0].height = Cm(0.65)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    set_repeat_header(table.rows[0])
    set_table_geometry(table, widths_cm)

    left = page_items[:25]
    right = page_items[25:50]
    for r in range(1, 26):
        set_row_cant_split(table.rows[r])
        table.rows[r].height = Cm(row_height_cm)
        table.rows[r].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for side, items in enumerate((left, right)):
            base = side * 4
            item = items[r - 1] if r - 1 < len(items) else None
            no = start_no + (r - 1) + side * 25
            vals = ["", "", "", ""]
            if item:
                vals[0] = str(no)
                vals[1] = "\n".join("★" for _ in range(int(item.get("tier", 1))))
                if is_answer:
                    vals[2] = item["word"]
                    vals[3] = item["meaning"]
                elif no % 2:
                    vals[2] = item["word"]
                else:
                    vals[2] = item["meaning"]
            for off, text in enumerate(vals):
                cell = table.rows[r].cells[base + off]
                cell.width = widths[base + off]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, top=0, start=30, bottom=0, end=30)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                if off == 1:
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    p.paragraph_format.line_spacing = 0.9
                    set_font(p.add_run(text), 7.0, name="Arial Unicode MS")
                else:
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    p.paragraph_format.line_spacing = 0.80
                    if off == 3:
                        size = fitted_font_size(text, 4.30, row_height_cm, 10.5)
                    elif off == 2:
                        size = fitted_font_size(text, 3.45, row_height_cm, 10.5)
                    else:
                        size = 10.0
                    set_font(p.add_run(text), size, name="함초롬바탕")
    return table


def add_footer_logo(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Cm(2.5))


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    data = json.loads(DATA.read_text(encoding="utf-8"))["BA"]
    data = sorted(enumerate(data), key=lambda x: (int(x[1].get("tier", 1)), x[1]["day"], x[0]))
    items = [x[1] for x in data]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(0.9)
    section.left_margin = Cm(1.1)
    section.right_margin = Cm(1.1)
    section.footer_distance = Cm(0.2)
    add_footer_logo(section)

    normal = doc.styles["Normal"]
    normal.font.name = "함초롬바탕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "함초롬바탕")
    normal.font.size = Pt(10.5)

    add_header(doc, "BASIC", "A")
    add_exam_table(doc, items[:50], 1, 0.86, False)
    doc.add_page_break()
    add_exam_table(doc, items[50:100], 51, 1.00, False)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
